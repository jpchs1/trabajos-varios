# -*- coding: utf-8 -*-
"""CV Jeniffer Mieres Contreras - Rediseno editorial ejecutivo.
Contenido identico al aprobado; solo cambia el tratamiento visual."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- PALETA ----------
INK    = RGBColor(0x16, 0x18, 0x1A)   # nombre, empresas, enfasis
BODY   = RGBColor(0x33, 0x37, 0x3B)   # cuerpo
META   = RGBColor(0x6C, 0x71, 0x76)   # fechas, cargos, instituciones
FAINT  = RGBColor(0xA8, 0xAD, 0xB2)   # separadores tipograficos
ACCENT = RGBColor(0x6E, 0x21, 0x30)   # oxblood - unico acento
HAIR   = "6E2130"                      # hairline del header

SERIF = "Cambria"      # primaria - editorial
SANS  = "Calibri"      # capa meta - etiquetas, fechas, cargos

# ---------- RETICULA ----------
M_TOP, M_BOT, M_LEFT, M_RIGHT = 1.5, 1.35, 1.8, 1.6
RAIL  = Cm(3.2)        # indentacion de la columna de contenido
BUL   = Cm(0.34)       # sangria del bullet

doc = Document()
s = doc.sections[0]
s.top_margin, s.bottom_margin = Cm(M_TOP), Cm(M_BOT)
s.left_margin, s.right_margin = Cm(M_LEFT), Cm(M_RIGHT)
FULL = s.page_width - s.left_margin - s.right_margin   # 17.6 cm

st = doc.styles["Normal"]
st.font.name = SERIF; st.font.size = Pt(9.5); st.font.color.rgb = BODY
st.element.rPr.rFonts.set(qn("w:eastAsia"), SERIF)
st.paragraph_format.space_before = Pt(0)
st.paragraph_format.space_after  = Pt(0)
st.paragraph_format.line_spacing = 1.16

def P(before=0, after=0, line=1.16, keep=False, indent=None, hang=None, tabs=()):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after, pf.line_spacing = Pt(before), Pt(after), line
    if keep: pf.keep_with_next = True
    if indent is not None: pf.left_indent = indent
    if hang   is not None: pf.first_line_indent = hang
    for pos, al in tabs: pf.tab_stops.add_tab_stop(pos, al)
    return p

def R(p, text, font=SERIF, size=9.5, bold=False, italic=False, color=BODY, track=None, caps=False):
    r = p.add_run(text)
    r.font.name = font; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    rPr = r._element.get_or_add_rPr()
    rPr.rFonts.set(qn("w:eastAsia"), font)
    rPr.rFonts.set(qn("w:cs"), font)
    if caps:
        e = OxmlElement("w:caps"); e.set(qn("w:val"), "1"); rPr.append(e)
    if track:
        e = OxmlElement("w:spacing"); e.set(qn("w:val"), str(track)); rPr.append(e)
    return r

def hairline(p, color=HAIR, sz=6, space=6):
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr"); bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), str(sz))
    bt.set(qn("w:space"), str(space)); bt.set(qn("w:color"), color)
    b.append(bt); pPr.append(b)

# --- etiqueta de seccion en el raíl izquierdo (comparte linea con el contenido) ---
def rail(label, before=15, after=0, keep=True, right_tab=False):
    tabs = [(RAIL, WD_TAB_ALIGNMENT.LEFT)]
    if right_tab: tabs.append((FULL, WD_TAB_ALIGNMENT.RIGHT))
    p = P(before=before, after=after, keep=keep, indent=RAIL, hang=-RAIL, tabs=tabs)
    R(p, label, font=SANS, size=8, bold=True, color=ACCENT, track=28, caps=True)
    p.add_run("\t")
    return p

def rail_bullet(label, before=15):
    """Etiqueta de seccion + primera vinieta en la misma linea."""
    p = P(before=before, indent=RAIL + BUL, hang=-(RAIL + BUL),
          tabs=[(RAIL, WD_TAB_ALIGNMENT.LEFT), (RAIL + BUL, WD_TAB_ALIGNMENT.LEFT)])
    R(p, label, font=SANS, size=8, bold=True, color=ACCENT, track=28, caps=True)
    p.add_run("\t")
    R(p, "\u2022", size=8.5, color=ACCENT)
    p.add_run("\t")
    return p

def body_par(before=0, after=0, line=1.16, keep=False):
    return P(before=before, after=after, line=line, keep=keep, indent=RAIL)

def bullet(pre="", key="", post="", after=2.7):
    p = P(after=after, indent=RAIL + BUL, hang=-BUL)
    R(p, "•", size=8.5, color=ACCENT)
    p.add_run("\t")
    if pre:  R(p, pre)
    if key:  R(p, key, bold=True, color=INK)
    if post: R(p, post)
    return p

def keywords(p, text, size=8.5):
    """Mismos terminos y separadores; los '·' pasan a gris claro."""
    for i, part in enumerate(text.split("  ·  ")):
        if i: R(p, "  ·  ", size=size, color=FAINT)
        R(p, part, size=size)

def job(company, role, dates, first=False, label=None):
    if first:
        p = rail(label, before=15, right_tab=True)
    else:
        p = P(before=9, keep=True, indent=RAIL,
              tabs=[(FULL, WD_TAB_ALIGNMENT.RIGHT)])
    R(p, company, size=10.5, bold=True, color=INK)
    p.add_run("\t")
    R(p, dates, font=SANS, size=8.5, color=META, track=14)
    q = P(before=2, after=4.5, keep=True, indent=RAIL)
    R(q, role, font=SANS, size=9, color=META, track=8)

def edu(title, inst, year, first=False, label=None):
    if first:
        p = rail(label, before=15, right_tab=True)
    else:
        p = P(before=8, keep=True, indent=RAIL,
              tabs=[(FULL, WD_TAB_ALIGNMENT.RIGHT)])
    R(p, title, size=9.5, bold=True, color=INK)
    p.add_run("\t")
    R(p, year, font=SANS, size=8.5, color=META, track=14)
    q = P(before=1, indent=RAIL)
    R(q, inst, font=SANS, size=8.5, color=META, track=8)

# ============================ ENCABEZADO ============================
p = P(after=2.5)
R(p, "JENIFFER MIERES CONTRERAS", size=21, color=INK, track=40)

p = P(before=0, after=3.5)
R(p, "Trabajadora Social", font=SANS, size=10.5, color=RGBColor(0x4A,0x4F,0x54), track=16)
R(p, "   |   ", font=SANS, size=10.5, color=ACCENT)
R(p, "Bienestar Laboral, Calidad de Vida y Gestión de Beneficios",
  font=SANS, size=10.5, color=RGBColor(0x4A,0x4F,0x54), track=16)

p = P(before=0, after=0)
for i, part in enumerate(["+56 9 4903 1682", "jeniffer.mieres@gmail.com",
                          "El Monte, Región Metropolitana"]):
    if i: R(p, "   ·   ", font=SANS, size=8.5, color=FAINT)
    R(p, part, font=SANS, size=8.5, color=META, track=12)
hairline(p, color=HAIR, sz=4, space=7)

# ============================ PERFIL ============================
p = rail("Perfil", before=16)
R(p, "Trabajadora Social titulada con distinción y más de 9 años de experiencia en bienestar laboral, gestión de "
     "beneficios y atención social directa, desarrollada en instituciones de educación superior y en empresa privada. "
     "Creé desde cero el Departamento de Bienestar Social de una organización con más de 200 colaboradores, "
     "administrando el Seguro Complementario de Salud, las licencias médicas y las cargas familiares de la dotación. "
     "Luego lideré la gestión de personas y bienestar en la consultora minera Bmining, a cargo de los programas de "
     "bienestar, convenios, plan anual de capacitación, clima laboral y de la matriz de riesgo y salud en el trabajo "
     "con seguimiento por KPI. Hoy me desempeño en la Universidad de Santiago de Chile, donde realizo evaluaciones "
     "socioeconómicas en los sistemas del Ministerio de Educación, gestiono beneficios ministeriales e internos y "
     "acompaño casos que requieren articular áreas internas con redes de apoyo. A la atención directa sumo una mirada "
     "de proceso: fui auditora interna del Sistema de Gestión Integrado (ISO 9001, 14001 y 45001). Cursando el "
     "Diplomado en Bienestar Organizacional y Estrategias de Diversidad e Inclusión (USACH) y certificada como agente "
     "Gatekeeper en prevención del suicidio.", size=9)
p.paragraph_format.line_spacing = 1.12

# ============================ ÁREAS DE EXPERTISE ============================
p = rail("Expertise", before=13)
p.paragraph_format.line_spacing = 1.2
keywords(p,
 "Creación, implementación y evaluación de programas de bienestar  ·  Gestión de beneficios internos y externos  ·  "
 "Evaluación socioeconómica e informes sociales  ·  Seguro Complementario de Salud  ·  "
 "Licencias médicas y cargas familiares  ·  Convenios institucionales y gestión de proveedores  ·  "
 "Gestión y seguimiento de casos sociales  ·  Orientación en salud y vivienda  ·  "
 "Coordinación de actividades y eventos  ·  Diseño y facilitación de talleres y capacitaciones  ·  "
 "Articulación con redes de apoyo  ·  Diversidad e inclusión  ·  Calidad de servicio  ·  "
 "Normativa laboral y gestión documental")

# ============================ EXPERIENCIA ============================
job("Universidad de Santiago de Chile", "Trabajadora Social · Departamento de Beneficios Estudiantiles",
    "Agosto 2023 – Actualidad", first=True, label="Experiencia")
bullet("Realizo la ", "evaluación socioeconómica",
       " de cada caso, revisando antecedentes y acreditaciones en los sistemas del Ministerio de Educación para "
       "definir con autonomía qué apoyo corresponde.")
bullet("Gestiono ", "beneficios ministeriales e internos",
       " de principio a fin: verificación de requisitos, preselección, seguimiento y registro documental que "
       "permite auditar el caso posteriormente.")
bullet("Atiendo al estudiantado de manera permanente y entrego ", "orientación social personalizada",
       ", acompañándolo durante todo el proceso de postulación.")
bullet("", "Coordino casos complejos",
       " con áreas internas de la universidad y con redes de apoyo externas cuando la necesidad excede el "
       "ámbito académico.")
bullet("Elaboro ", "informes sociales",
       " y mantengo la trazabilidad administrativa de la cartera de casos a mi cargo.")

job("Bmining · Desarrollo e Innovación para la Minería SpA",
    "Líder en Gestión de Personas · Consultora para la industria minera", "Octubre 2018 – Febrero 2023")
bullet("Estuve a cargo de los ", "programas de bienestar y beneficios",
       " de los colaboradores, desde la definición de la oferta hasta su implementación y seguimiento, con foco "
       "en el clima laboral y en el día a día de la empresa.")
bullet("Gestioné ", "convenios y beneficios orientados a la calidad de vida laboral",
       ", administrando la relación con instituciones y proveedores externos.")
bullet("Estuve a cargo del área de Recursos Humanos en ", "selección e inducción",
       ", y lideré el plan anual de capacitaciones internas.")
bullet("Tuve a cargo la ", "matriz de riesgo y salud en el trabajo",
       " junto al prevencionista externo, con seguimiento por KPI para dar cumplimiento a la certificación.")
bullet("Impulsé el ", "programa de cuidado del medio ambiente",
       " con actividades semanales y seguimiento, con participación de toda la organización desde la gerencia "
       "general en adelante.")
bullet("Organicé ", "actividades corporativas de bienestar e integración",
       ", articulando a las distintas áreas para asegurar convocatoria y buena ejecución.")
bullet("Asumí durante un año el ", "Sistema de Gestión Integrado como auditora interna",
       " en las certificaciones ISO 9001:2015 (Calidad), ISO 14001:2015 (Medio Ambiente) e ISO 45001:2018 "
       "(Seguridad y Salud en el Trabajo).")

job("Fundación Fondo Esperanza", "Monitora · Grupo Emprendedores", "Octubre 2017 – Septiembre 2018")
bullet("Diseñé y dicté ", "capacitaciones",
       " para emprendedores que estaban partiendo o buscando hacer crecer su negocio, con metodologías "
       "prácticas y grupales.")
bullet("", "Acompañé la gestión de sus proyectos",
       ", revisando con cada participante qué reforzar para alcanzar sus objetivos.")
bullet("Orienté a los emprendedores en materia de ", "vivienda",
       ", articulando con las redes de la comuna, y los vinculé con recursos disponibles que muchas veces no "
       "sabían que tenían a su alcance.")

job("Ferretería San Francisco", "Encargada · Departamento de Bienestar Social",
    "Diciembre 2016 – Septiembre 2017")
bullet("", "Creé desde cero el área de Bienestar Social",
       " para una dotación de más de 200 colaboradores, definiendo la planificación de actividades y beneficios.")
bullet("Administré el ", "Seguro Complementario de Salud",
       " de la dotación, apoyando a los trabajadores en el acceso, uso y tramitación de reembolsos.")
bullet("Gestioné ", "licencias médicas ante la Caja de Compensación Los Andes",
       " y la incorporación de cargas familiares de los trabajadores.")
bullet("Realicé ", "visitas semanales a la segunda sucursal",
       ", atendiendo en terreno las necesidades de beneficios internos y externos, principalmente en "
       "vivienda y salud.")
bullet("Gestioné ", "convenios con instituciones externas", " para mejorar la calidad de vida en el trabajo.")
bullet("Administré vacaciones, carpetas personales y documentación laboral, y apoyé los ",
       "procesos administrativos de Recursos Humanos", ".")

# ============================ FORMACIÓN ============================
edu("Diplomado en Bienestar Organizacional y Estrategias de Diversidad e Inclusión",
    "Universidad de Santiago de Chile", "En curso", first=True, label="Formación")
edu("Postítulo en Trabajo Social en Niñez, Adolescencia y Familia en el Contexto Judicial",
    "Universidad Andrés Bello", "2016")
edu("Trabajadora Social · Título profesional con distinción",
    "Instituto Profesional AIEP", "2012")

# ============================ CERTIFICACIONES ============================
p = rail_bullet("Certificaciones", before=15)
R(p, "Certificación Gatekeepers", bold=True, color=INK)
R(p, " — agentes comunitarios en prevención del suicidio (2026).")
bullet("", "Interpretación auditor interno",
       " — TÜV Rheinland (2019). Respalda la auditoría del Sistema de Gestión Integrado ISO 9001 / 14001 / 45001.")
bullet("", "Ley 21.327, Modernización de la Dirección del Trabajo", " — Bicentenario OTEC (2021).")
bullet("", "Cálculo de finiquito", " — Consultores y Asesores en Capacitación Chile Ltda. (2022).")

# ============================ COMPETENCIAS ============================
p = rail("Competencias", before=15)
p.paragraph_format.line_spacing = 1.2
keywords(p,
 "Vocación y orientación al servicio  ·  Excelente trato al usuario y altos estándares de calidad de atención  ·  "
 "Comunicación efectiva  ·  Trabajo en equipo y coordinación transversal con múltiples unidades  ·  "
 "Autonomía y resolución de situaciones  ·  Flexibilidad y adaptación al cambio  ·  Manejo de conflictos  ·  "
 "Proactividad y aporte al clima laboral  ·  Rigurosidad administrativa y normativa")

# ============================ HERRAMIENTAS ============================
p = rail("Herramientas", before=15)
R(p, "Microsoft Office (Word, Excel y PowerPoint) nivel intermedio; conocimiento básico de BUK, SAP y Talana; "
     "plataformas institucionales del Ministerio de Educación y gestión documental.")

p = rail("Disponibilidad", before=10)
R(p, "Movilización propia y licencia clase B; disponibilidad para desempeñarse en Campus San Joaquín y en "
     "otros campus de la Región Metropolitana.")

out = "/home/user/trabajos-varios/cv-trabajador-social-uc/CV_Jeniffer_Mieres_Contreras_Trabajadora_Social_UC.docx"
doc.save(out)
print("OK ->", out)
