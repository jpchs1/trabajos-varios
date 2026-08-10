# -*- coding: utf-8 -*-
"""Tres variantes de diseno sobre el mismo contenido (cv_content.py)."""
import sys, copy
sys.path.insert(0, "/tmp/claude-0/-home-user-trabajos-varios/08e63ebe-9bab-5b5e-bcb3-d86e109b9bf1/scratchpad/cvwork")
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import cv_content as C

OUT = "/home/user/trabajos-varios/cv-trabajador-social-uc/"

# --- knobs de densidad (los ajusta el afinador automatico) ---
K = dict(body=0.0, line=0.0, bul=0.0, sec=0.0, job=0.0, par=0.0, mar=0.0)
def kb(v): return v + K["body"]          # tamano de cuerpo
def kl(v): return v + K["line"]          # interlineado
def kbu(v): return max(1.0, v + K["bul"])# aire entre vinietas
def ks(v): return max(6.0, v + K["sec"]) # aire antes de seccion
def kj(v): return max(5.0, v + K["job"]) # aire antes de cargo
def kp(v): return max(2.0, v + K["par"]) # aire entre parrafos de perfil
def km(v): return max(0.7, v + K["mar"]) # margenes vert.

# ------------------------------------------------------------------ utilidades
def el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items(): e.set(qn("w:" + k), str(v))
    return e

def hyphenate(doc, zone=510):
    s = doc.settings.element
    for tag, val in (("w:autoHyphenation","1"), ("w:consecutiveHyphenLimit","2"),
                     ("w:doNotHyphenateCaps","1"), ("w:hyphenationZone", str(zone))):
        s.append(el(tag, val=val))

def base_style(doc, font, size, color, line):
    st = doc.styles["Normal"]
    st.font.name = font; st.font.size = Pt(size); st.font.color.rgb = color
    rPr = st.element.get_or_add_rPr()
    rPr.rFonts.set(qn("w:eastAsia"), font)
    rPr.append(el("w:lang", val="es-CL"))
    pf = st.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = line

def R(p, text, font, size, bold=False, color=None, track=None, caps=False, italic=False):
    r = p.add_run(text)
    r.font.name = font; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color
    rp = r._element.get_or_add_rPr()
    rp.rFonts.set(qn("w:eastAsia"), font); rp.rFonts.set(qn("w:cs"), font)
    if caps:  rp.append(el("w:caps", val="1"))
    if track: rp.append(el("w:spacing", val=track))
    return r

def border(p, edge, color, sz, space=4):
    pPr = p._p.get_or_add_pPr()
    b = pPr.find(qn("w:pBdr"))
    if b is None:
        b = OxmlElement("w:pBdr"); pPr.append(b)
    b.append(el("w:" + edge, val="single", sz=sz, space=space, color=color))

def shade(p, fill):
    p._p.get_or_add_pPr().append(el("w:shd", val="clear", fill=fill))

def cell_shade(c, fill):
    c._tc.get_or_add_tcPr().append(el("w:shd", val="clear", fill=fill))

def cell_pad(c, top=0, start=0, bottom=0, end=0):
    m = OxmlElement("w:tcMar")
    for tag, v in (("top",top), ("left",start), ("bottom",bottom), ("right",end)):
        m.append(el("w:" + tag, w=int(v*567), type="dxa"))
    c._tc.get_or_add_tcPr().append(m)

def keep_block(paras):
    """Un cargo nunca se parte: si no cabe, arranca en la pagina siguiente."""
    paras = [x for x in paras if x is not None]
    for i, x in enumerate(paras):
        x.paragraph_format.keep_together  = True
        x.paragraph_format.keep_with_next = (i < len(paras) - 1)

def no_hyph(p):
    p._p.get_or_add_pPr().append(el("w:suppressAutoHyphens", val="1"))

def fixed_table(t, widths):
    """Anchos de columna reales: tblW + tblGrid + ancho por celda."""
    t.autofit = False
    tbl = t._tbl
    tblPr = tbl.tblPr
    tblPr.append(el("w:tblLayout", type="fixed"))
    total = sum(int(w.twips) for w in widths)
    tblPr.append(el("w:tblW", w=total, type="dxa"))
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None: tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths: grid.append(el("w:gridCol", w=int(w.twips)))
    tbl.insert(list(tbl).index(tblPr) + 1, grid)
    for row in t.rows:
        row._tr.append(el("w:trPr")) if row._tr.find(qn("w:trPr")) is None else None
        for i, w in enumerate(widths):
            c = row.cells[i]
            c.width = w
            c._tc.get_or_add_tcPr().append(el("w:tcW", w=int(w.twips), type="dxa"))

# =============================================================================
# VARIANTE 1 — SWISS  (sans puro, sin color, reglas cortas, maximo blanco)
# =============================================================================
def swiss():
    F = "Calibri"
    INK  = RGBColor(0x0D,0x0D,0x0D); BODY = RGBColor(0x3A,0x3E,0x42)
    META = RGBColor(0x76,0x7B,0x80); FAINT= RGBColor(0xB0,0xB4,0xB8)
    doc = Document(); s = doc.sections[0]
    s.top_margin, s.bottom_margin = Cm(km(1.5)), Cm(km(1.3))
    s.left_margin, s.right_margin = Cm(1.6), Cm(1.45)
    W = s.page_width - s.left_margin - s.right_margin
    base_style(doc, F, kb(9.5), BODY, kl(1.18)); hyphenate(doc)

    def P(before=0, after=0, line=None, keep=False, indent=None, hang=None, tabs=(),
          just=False, rind=None):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.space_before, pf.space_after = Pt(before), Pt(after)
        pf.line_spacing = kl(1.18) if line is None else line
        if keep: pf.keep_with_next = True
        if indent is not None: pf.left_indent = indent
        if hang is not None: pf.first_line_indent = hang
        if rind is not None: pf.right_indent = rind
        for pos, al in tabs: pf.tab_stops.add_tab_stop(pos, al)
        if just: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def head(txt, before=None):
        before = ks(17) if before is None else ks(before)
        # regla corta y gruesa sobre el titulo: acento grafico suizo
        r = P(before=before, after=0, line=1, rind=W - Cm(2.2))
        border(r, "top", "0D0D0D", 18, space=0)
        R(r, "", F, 2)
        h = P(before=5, after=6, keep=True)
        R(h, txt.upper(), F, 8, bold=True, color=INK, track=38)

    def bullet(pre, key, post, after=None):
        after = kbu(4.2) if after is None else after
        p = P(after=after, indent=Cm(0.42), hang=Cm(-0.42), just=True)
        R(p, "—", F, 9, color=FAINT); p.add_run("\t")
        if pre:  R(p, pre, F, 9.5)
        if key:  R(p, key, F, 9.5, bold=True, color=INK)
        if post: R(p, post, F, 9.5)
        return p

    # encabezado
    p = P(after=3); R(p, C.NOMBRE, F, 25, bold=True, color=INK, track=-4)
    p = P(after=7); R(p, (C.TITULO + "   ·   " + C.CLAIM).upper(), F, 8, color=META, track=34)
    p = P(after=0, tabs=[(W, WD_TAB_ALIGNMENT.RIGHT)])
    R(p, C.FONO + "   ·   " + C.MAIL, F, 9, color=BODY)
    R(p, "\t" + C.CIUDAD, F, 9, color=META)
    border(p, "bottom", "0D0D0D", 8, space=8)

    head("Perfil", before=15)
    for i, t in enumerate(C.PERFIL):
        q = P(before=0 if i == 0 else kp(5), just=True); R(q, t, F, kb(9.5))

    head("Áreas de expertise")
    p = P(line=kl(1.34)); no_hyph(p)
    for i, t in enumerate(C.EXPERTISE):
        if i: R(p, "   ·   ", F, 9, color=FAINT)
        R(p, t, F, 9)

    head("Experiencia profesional")
    for j, job in enumerate(C.JOBS):
        p = P(before=0 if j == 0 else kj(12), keep=True, tabs=[(W, WD_TAB_ALIGNMENT.RIGHT)])
        if j == 2: p.paragraph_format.page_break_before = True
        R(p, job["empresa"], F, kb(11), bold=True, color=INK)
        R(p, "\t" + job["fechas"], F, 7.4, color=META, track=24, caps=True)
        q = P(before=2, after=5, keep=True); R(q, job["cargo"], F, 9, color=META, track=8, italic=True)
        keep_block([p, q] + [bullet(*b) for b in job["bullets"]])

    head("Formación académica")
    for i, (t, inst, y) in enumerate(C.EDU):
        p = P(before=0 if i == 0 else 8, keep=True, tabs=[(W, WD_TAB_ALIGNMENT.RIGHT)])
        R(p, t, F, 9.5, bold=True, color=INK); R(p, "\t" + y, F, 7.4, color=META, track=24, caps=True)
        q = P(before=1.5); R(q, inst, F, 9, color=META); keep_block([p, q])

    head("Cursos y certificaciones")
    for k, rest in C.CERTS: bullet("", k, rest)

    head("Competencias")
    p = P(line=kl(1.34)); no_hyph(p)
    for i, t in enumerate(C.COMPETENCIAS):
        if i: R(p, "   ·   ", F, 9, color=FAINT)
        R(p, t, F, 9)

    head("Herramientas y disponibilidad")
    bullet("", "Herramientas: ", C.HERRAMIENTAS)
    bullet("", "Disponibilidad: ", C.DISPONIBILIDAD)

    f = OUT + "CV_Jeniffer_Mieres_01_Minimal.docx"; doc.save(f); return f

# =============================================================================
# VARIANTE 2 — BANDA  (encabezado en bloque de color, cuerpo limpio)
# =============================================================================
def banda():
    DISP, TXT = "Cambria", "Calibri"
    DEEP = "13323C"
    INK  = RGBColor(0x15,0x1A,0x1C); BODY = RGBColor(0x35,0x3A,0x3D)
    META = RGBColor(0x6E,0x74,0x78); FAINT= RGBColor(0xAE,0xB3,0xB6)
    ACC  = RGBColor(0x13,0x32,0x3C); WHITE= RGBColor(0xFF,0xFF,0xFF)
    PALE = RGBColor(0xC5,0xD2,0xD6)
    doc = Document(); s = doc.sections[0]
    s.top_margin, s.bottom_margin = Cm(km(1.15)), Cm(km(1.25))
    s.left_margin, s.right_margin = Cm(1.6), Cm(1.45)
    W = s.page_width - s.left_margin - s.right_margin
    base_style(doc, TXT, kb(10), BODY, kl(1.16)); hyphenate(doc)

    def P(before=0, after=0, line=None, keep=False, indent=None, hang=None, tabs=(),
          just=False, lind=None, rind=None):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.space_before, pf.space_after = Pt(before), Pt(after)
        pf.line_spacing = kl(1.16) if line is None else line
        if keep: pf.keep_with_next = True
        if indent is not None: pf.left_indent = indent
        if lind   is not None: pf.left_indent = lind
        if hang   is not None: pf.first_line_indent = hang
        if rind   is not None: pf.right_indent = rind
        for pos, al in tabs: pf.tab_stops.add_tab_stop(pos, al)
        if just: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def head(txt, before=15):
        h = P(before=ks(before), after=1, keep=True)
        R(h, txt.upper(), TXT, 8.5, bold=True, color=ACC, track=48)
        r = P(before=1.5, after=6, line=1, keep=True, rind=W - Cm(1.4))
        border(r, "bottom", DEEP, 16, space=0); R(r, "", TXT, 2)

    def bullet(pre, key, post, after=None):
        after = kbu(4.2) if after is None else after
        p = P(after=after, indent=Cm(0.40), hang=Cm(-0.40), just=True)
        R(p, "▪", TXT, 7, color=ACC); p.add_run("\t")
        if pre:  R(p, pre, TXT, 10)
        if key:  R(p, key, TXT, 10, bold=True, color=INK)
        if post: R(p, post, TXT, 10)
        return p

    # ---- banda superior a sangre ----
    BLEED_L, BLEED_R = Cm(-1.6), Cm(-1.45)
    b1 = P(before=5, after=1, lind=BLEED_L, rind=BLEED_R)
    shade(b1, DEEP); R(b1, "  " + C.NOMBRE_T, DISP, 20, color=WHITE, track=10)
    b2 = P(after=2, lind=BLEED_L, rind=BLEED_R)
    shade(b2, DEEP); R(b2, "  " + C.TITULO.upper() + "   |   " + C.CLAIM.upper(),
                       TXT, 8, bold=True, color=PALE, track=34)
    b3 = P(after=0, lind=BLEED_L, rind=BLEED_R)
    shade(b3, DEEP)
    R(b3, "  " + C.FONO + "   ·   " + C.MAIL + "   ·   " + C.CIUDAD, TXT, 8.5, color=PALE)
    b4 = P(after=11, line=1, lind=BLEED_L, rind=BLEED_R)
    shade(b4, DEEP); R(b4, "", TXT, 4)

    head("Perfil profesional", before=0)
    for i, t in enumerate(C.PERFIL):
        q = P(before=0 if i == 0 else kp(5), just=True); R(q, t, TXT, kb(10))

    head("Áreas de expertise")
    p = P(line=kl(1.32)); no_hyph(p)
    for i, t in enumerate(C.EXPERTISE):
        if i: R(p, "   ·   ", TXT, 9.5, color=FAINT)
        R(p, t, TXT, 9.5)

    head("Experiencia profesional")
    for j, job in enumerate(C.JOBS):
        p = P(before=0 if j == 0 else kj(12), keep=True, tabs=[(W, WD_TAB_ALIGNMENT.RIGHT)])
        if j == 2: p.paragraph_format.page_break_before = True
        R(p, job["empresa"], DISP, kb(10.5), bold=True, color=INK)
        R(p, "\t" + job["fechas"], TXT, 7.6, color=META, track=22, caps=True)
        q = P(before=2, after=5, keep=True); R(q, job["cargo"], DISP, 9.5, color=META, track=4, italic=True)
        keep_block([p, q] + [bullet(*b) for b in job["bullets"]])

    head("Formación académica")
    for i, (t, inst, y) in enumerate(C.EDU):
        p = P(before=0 if i == 0 else 8, keep=True, tabs=[(W, WD_TAB_ALIGNMENT.RIGHT)])
        R(p, t, DISP, 10, bold=True, color=INK); R(p, "\t" + y, TXT, 7.6, color=META, track=22, caps=True)
        q = P(before=1.5); R(q, inst, TXT, 9.5, color=META); keep_block([p, q])

    head("Cursos y certificaciones")
    for k, rest in C.CERTS: bullet("", k, rest)

    head("Competencias")
    p = P(line=kl(1.32)); no_hyph(p)
    for i, t in enumerate(C.COMPETENCIAS):
        if i: R(p, "   ·   ", TXT, 9.5, color=FAINT)
        R(p, t, TXT, 9.5)

    head("Herramientas y disponibilidad")
    bullet("", "Herramientas: ", C.HERRAMIENTAS)
    bullet("", "Disponibilidad: ", C.DISPONIBILIDAD)

    f = OUT + "CV_Jeniffer_Mieres_02_Banda.docx"; doc.save(f); return f

# =============================================================================
# VARIANTE 3 — SIDEBAR  (dos columnas, barra lateral tramada)
# =============================================================================
def sidebar():
    F = "Calibri"
    DEEP = "1B3A4B"
    INK  = RGBColor(0x15,0x1A,0x1C); BODY = RGBColor(0x35,0x3A,0x3D)
    META = RGBColor(0x6E,0x74,0x78); ACC  = RGBColor(0x1B,0x3A,0x4B)
    WHITE= RGBColor(0xFF,0xFF,0xFF); PALE = RGBColor(0xC7,0xD4,0xDA)
    doc = Document(); s = doc.sections[0]
    s.top_margin, s.bottom_margin = Cm(km(1.1)), Cm(km(1.1))
    s.left_margin, s.right_margin = Cm(1.5), Cm(1.4)
    SIDE, MAIN = Cm(5.4), Cm(12.7)
    base_style(doc, F, kb(9.5), BODY, kl(1.15)); hyphenate(doc, zone=397)

    def par(cont, before=0, after=0, line=None, keep=False, indent=None, hang=None,
            just=False, lind=None, rind=None, brk=False):
        p = cont.add_paragraph(); pf = p.paragraph_format
        pf.space_before, pf.space_after = Pt(before), Pt(after)
        pf.line_spacing = kl(1.15) if line is None else line
        if keep: pf.keep_with_next = True
        if indent is not None: pf.left_indent = indent
        if lind   is not None: pf.left_indent = lind
        if hang   is not None: pf.first_line_indent = hang
        if rind   is not None: pf.right_indent = rind
        if just:  p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if brk:   pf.page_break_before = True
        return p

    def head(cont, txt, before=13, first=False):
        h = par(cont, before=0 if first else ks(before), after=1, keep=True)
        R(h, txt.upper(), F, 8, bold=True, color=ACC, track=44)
        r = par(cont, before=1.5, after=5, line=1, keep=True)
        border(r, "bottom", DEEP, 10, space=0); R(r, "", F, 2)

    EX = [0.0]
    def bullet(cont, pre, key, post):
        p = par(cont, after=kbu(4.0) + EX[0], indent=Cm(0.36), hang=Cm(-0.36), just=True)
        R(p, "\u25aa", F, 6.5, color=ACC); p.add_run("\t")
        if pre:  R(p, pre, F, kb(9.5))
        if key:  R(p, key, F, kb(9.5), bold=True, color=INK)
        if post: R(p, post, F, kb(9.5))
        return p

    def tag(cont, txt):
        p = par(cont, after=kbu(2.8) + EX[0], indent=Cm(0.3), hang=Cm(-0.3), line=kl(1.10))
        no_hyph(p); R(p, "\u25aa", F, 6.5, color=ACC); p.add_run("\t")
        R(p, txt, F, kb(8.8))

    def banda_sup(brk=False):
        b1 = par(doc, before=0 if brk else 6, after=2, lind=Cm(-1.5), rind=Cm(-1.4), brk=brk)
        shade(b1, DEEP); R(b1, "  " + C.NOMBRE_T, F, kb(23), bold=True, color=WHITE, track=6)
        b2 = par(doc, after=2, lind=Cm(-1.5), rind=Cm(-1.4))
        shade(b2, DEEP)
        R(b2, "  " + C.TITULO.upper() + "   |   " + C.CLAIM.upper(), F, kb(8.5),
          bold=True, color=PALE, track=36)
        b3 = par(doc, after=0, lind=Cm(-1.5), rind=Cm(-1.4))
        shade(b3, DEEP)
        R(b3, "  " + C.FONO + "   \u00b7   " + C.MAIL + "   \u00b7   " + C.CIUDAD,
          F, kb(9), color=PALE)
        b4 = par(doc, after=8, line=1, lind=Cm(-1.5), rind=Cm(-1.4))
        shade(b4, DEEP); R(b4, "", F, 5)

    def make_table():
        t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
        fixed_table(t, [SIDE, MAIN])
        L, Rc = t.rows[0].cells[0], t.rows[0].cells[1]
        cell_shade(L, "EFF3F5")
        tb = OxmlElement("w:tcBorders")
        tb.append(el("w:start", val="single", sz=18, space=0, color=DEEP))
        L._tc.get_or_add_tcPr().append(tb)
        cell_pad(L, 0.32, 0.32, 0.32, 0.32); cell_pad(Rc, 0.32, 0.5, 0.32, 0)
        for c in (L, Rc):
            c.paragraphs[0].text = ""
            c._tc.get_or_add_tcPr().append(el("w:vAlign", val="top"))
        return L, Rc

    def job_block(cont, job):
        p = par(cont, keep=True); R(p, job["empresa"], F, kb(10.5), bold=True, color=INK)
        q = par(cont, before=1.5, keep=True); R(q, job["cargo"], F, kb(9), color=META)
        d = par(cont, before=0.5, after=4, keep=True)
        R(d, job["fechas"], F, 7.6, bold=True, color=ACC, track=20, caps=True)
        keep_block([p, q, d] + [bullet(cont, *b) for b in job["bullets"]])
        return p

    # ---------------- PAGINA 1 ----------------
    banda_sup()
    L, Rc = make_table()
    head(L, "Áreas de expertise", first=True)
    for x in C.EXPERTISE: tag(L, x)
    head(L, "Competencias")
    for x in C.COMPETENCIAS: tag(L, x)

    head(Rc, "Perfil profesional", first=True)
    for i, x in enumerate(C.PERFIL):
        q = par(Rc, before=0 if i == 0 else kp(4.5), just=True); R(q, x, F, kb(9.5))
    head(Rc, "Experiencia profesional")
    for j, job in enumerate(C.JOBS[:2]):
        p = job_block(Rc, job)
        if j: p.paragraph_format.space_before = Pt(kj(10.5))

    # ---------------- PAGINA 2 ----------------
    EX[0] = 3.2                      # pagina 2 respira mas: equilibra el largo
    banda_sup(brk=True)
    L2, R2 = make_table()
    head(L2, "Herramientas", first=True)
    p = par(L2, line=kl(1.10)); no_hyph(p); R(p, C.HERRAMIENTAS, F, kb(8.8))
    head(L2, "Disponibilidad", before=20)
    p = par(L2, line=kl(1.10)); no_hyph(p); R(p, C.DISPONIBILIDAD, F, kb(8.8))
    head(L2, "Cursos y certificaciones", before=20)
    for k, rest in C.CERTS: tag(L2, k + rest)

    head(R2, "Experiencia profesional (cont.)", first=True)
    for j, job in enumerate(C.JOBS[2:]):
        p = job_block(R2, job)
        if j: p.paragraph_format.space_before = Pt(kj(10.5) + 8)
    head(R2, "Formación académica", before=22)
    for i, (x, inst, y) in enumerate(C.EDU):
        p = par(R2, before=0 if i == 0 else 13, keep=True)
        R(p, x, F, kb(9.5), bold=True, color=INK)
        q = par(R2, before=1); R(q, inst + "   \u00b7   " + y, F, kb(9), color=META)
        keep_block([p, q])

    f = OUT + "CV_Jeniffer_Mieres_03_Sidebar.docx"; doc.save(f); return f


# =============================================================================
# VARIANTE 4 — EDITORIAL (rail de etiquetas, serif, acento oxblood)
# =============================================================================
def rail():
    SERIF, SANS = "Cambria", "Calibri"
    INK = RGBColor(0x16,0x18,0x1A); BODY = RGBColor(0x33,0x37,0x3B)
    META = RGBColor(0x6C,0x71,0x76); FAINT = RGBColor(0xA8,0xAD,0xB2)
    ACC = RGBColor(0x6E,0x21,0x30); HAIR = "6E2130"
    doc = Document(); s = doc.sections[0]
    s.top_margin, s.bottom_margin = Cm(km(1.25)), Cm(km(1.15))
    s.left_margin, s.right_margin = Cm(1.7), Cm(1.5)
    W = s.page_width - s.left_margin - s.right_margin
    RAILW, BUL = Cm(2.9), Cm(0.34)
    base_style(doc, SERIF, kb(9.5), BODY, kl(1.21)); hyphenate(doc, zone=510)

    def P(before=0, after=0, line=None, keep=False, indent=None, hang=None, tabs=(), just=False, brk=False):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.space_before, pf.space_after = Pt(before), Pt(after)
        pf.line_spacing = kl(1.21) if line is None else line
        if keep: pf.keep_with_next = True
        if indent is not None: pf.left_indent = indent
        if hang is not None: pf.first_line_indent = hang
        for pos, al in tabs: pf.tab_stops.add_tab_stop(pos, al)
        if just: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if brk: pf.page_break_before = True
        return p

    def railp(label, before, right_tab=False, just=True, brk=False):
        tabs = [(RAILW, WD_TAB_ALIGNMENT.LEFT)]
        if right_tab: tabs.append((W, WD_TAB_ALIGNMENT.RIGHT))
        p = P(before=ks(before), keep=True, indent=RAILW, hang=-RAILW, tabs=tabs, just=just, brk=brk)
        R(p, label, SANS, 7.5, bold=True, color=ACC, track=20, caps=True); p.add_run("\t")
        return p

    def bullet(pre, key, post):
        p = P(after=kbu(4.6), indent=RAILW+BUL, hang=-BUL, just=True)
        R(p, "\u2022", SERIF, 8.5, color=ACC); p.add_run("\t")
        if pre: R(p, pre, SERIF, kb(9.5))
        if key: R(p, key, SERIF, kb(9.5), bold=True, color=INK)
        if post: R(p, post, SERIF, kb(9.5))
        return p

    p = P(after=2); R(p, C.NOMBRE, SERIF, 20, color=INK, track=36)
    p = P(after=3)
    R(p, C.TITULO, SANS, 10, color=RGBColor(0x4A,0x4F,0x54), track=14)
    R(p, "   |   ", SANS, 10, color=ACC)
    R(p, C.CLAIM, SANS, 10, color=RGBColor(0x4A,0x4F,0x54), track=14)
    p = P(after=0)
    R(p, C.FONO + "   \u00b7   " + C.MAIL + "   \u00b7   " + C.CIUDAD, SANS, 8.5, color=META, track=10)
    border(p, "bottom", HAIR, 4, space=7)

    p = railp("Perfil", 14); R(p, C.PERFIL[0], SERIF, kb(9.5))
    for t in C.PERFIL[1:]:
        q = P(before=kp(6), indent=RAILW, just=True); R(q, t, SERIF, kb(9.5))

    p = railp("Expertise", 13, just=False); no_hyph(p)
    p.paragraph_format.line_spacing = kl(1.3)
    for i, t in enumerate(C.EXPERTISE):
        if i: R(p, "  \u00b7  ", SERIF, 8.5, color=FAINT)
        R(p, t, SERIF, 8.5)

    def job(j, first=False, label=None, brk=False):
        if first:
            p = railp(label, 14, right_tab=True, just=False, brk=brk)
        else:
            p = P(before=kj(12), keep=True, indent=RAILW,
                  tabs=[(W, WD_TAB_ALIGNMENT.RIGHT)], brk=brk)
        R(p, j["empresa"], SERIF, kb(10.5), bold=True, color=INK); p.add_run("\t")
        R(p, j["fechas"], SANS, 7.6, color=META, track=20, caps=True)
        q = P(before=1.5, after=3.5, keep=True, indent=RAILW)
        R(q, j["cargo"], SANS, 8.8, color=META, track=6, italic=True)
        keep_block([p, q] + [bullet(*b) for b in j["bullets"]])

    job(C.JOBS[0], first=True, label="Experiencia")
    job(C.JOBS[1])
    job(C.JOBS[2], brk=True)
    job(C.JOBS[3])

    def edu(t, inst, y, first=False):
        if first:
            p = railp("Formaci\u00f3n", 14, right_tab=True, just=False)
        else:
            p = P(before=6.5, keep=True, indent=RAILW, tabs=[(W, WD_TAB_ALIGNMENT.RIGHT)])
        R(p, t, SERIF, kb(9.5), bold=True, color=INK); p.add_run("\t")
        R(p, y, SANS, 7.6, color=META, track=20, caps=True)
        q = P(before=1, indent=RAILW); R(q, inst, SANS, 8.5, color=META, italic=True)
        keep_block([p, q])
    for i, (t, inst, y) in enumerate(C.EDU): edu(t, inst, y, first=(i == 0))

    p = railp("Certificaciones", 13)
    k, r = C.CERTS[0]
    R(p, k, SERIF, kb(9.5), bold=True, color=INK); R(p, r, SERIF, kb(9.5))
    for k, r in C.CERTS[1:]: bullet("", k, r)

    p = railp("Competencias", 13, just=False); no_hyph(p)
    p.paragraph_format.line_spacing = kl(1.3)
    for i, t in enumerate(C.COMPETENCIAS):
        if i: R(p, "  \u00b7  ", SERIF, 8.5, color=FAINT)
        R(p, t, SERIF, 8.5)
    p = railp("Herramientas", 12); R(p, C.HERRAMIENTAS, SERIF, kb(9))
    p = railp("Disponibilidad", 9); R(p, C.DISPONIBILIDAD, SERIF, kb(9))

    f = OUT + "CV_Jeniffer_Mieres_04_Editorial.docx"; doc.save(f); return f

# =============================================================================
# VARIANTE 5 — CLASICO (centrado, serif, doble filete, acento bronce)
# =============================================================================
def centro():
    SERIF, SANS = "Cambria", "Calibri"
    INK = RGBColor(0x1A,0x1A,0x1A); BODY = RGBColor(0x35,0x38,0x3B)
    META = RGBColor(0x70,0x74,0x78); ACC = RGBColor(0x7A,0x5C,0x2E); HAIR = "7A5C2E"
    doc = Document(); s = doc.sections[0]
    s.top_margin, s.bottom_margin = Cm(km(1.3)), Cm(km(1.2))
    s.left_margin, s.right_margin = Cm(1.9), Cm(1.9)
    W = s.page_width - s.left_margin - s.right_margin
    base_style(doc, SERIF, kb(9.5), BODY, kl(1.18)); hyphenate(doc, zone=510)

    def P(before=0, after=0, line=None, keep=False, indent=None, hang=None, tabs=(),
          just=False, center=False, brk=False, lind=None, rind=None):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.space_before, pf.space_after = Pt(before), Pt(after)
        pf.line_spacing = kl(1.18) if line is None else line
        if keep: pf.keep_with_next = True
        if indent is not None: pf.left_indent = indent
        if lind is not None: pf.left_indent = lind
        if rind is not None: pf.right_indent = rind
        if hang is not None: pf.first_line_indent = hang
        for pos, al in tabs: pf.tab_stops.add_tab_stop(pos, al)
        if just: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if brk: pf.page_break_before = True
        return p

    def head(t, before=14, brk=False):
        h = P(before=ks(before), after=1, keep=True, center=True, brk=brk)
        R(h, t.upper(), SANS, 8, bold=True, color=ACC, track=44)
        from docx.shared import Emu
        side = Emu(int((W - Cm(2.4)) / 2))
        r = P(before=1.5, after=6, line=1, keep=True, lind=side, rind=side)
        border(r, "bottom", HAIR, 8, space=0); R(r, "", SANS, 2)

    def bullet(pre, key, post):
        p = P(after=kbu(3.4), indent=Cm(0.4), hang=Cm(-0.4), just=True)
        R(p, "\u00b7", SERIF, 10, bold=True, color=ACC); p.add_run("\t")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(0.4), WD_TAB_ALIGNMENT.LEFT)
        if pre: R(p, pre, SERIF, kb(9.5))
        if key: R(p, key, SERIF, kb(9.5), bold=True, color=INK)
        if post: R(p, post, SERIF, kb(9.5))
        return p

    p = P(after=3, center=True); R(p, C.NOMBRE, SERIF, 19, color=INK, track=52, caps=True)
    p = P(after=2.5, center=True)
    R(p, C.TITULO.upper(), SANS, 8, bold=True, color=META, track=36)
    R(p, "   \u00b7   ", SANS, 8, color=ACC)
    R(p, C.CLAIM.upper(), SANS, 8, bold=True, color=META, track=36)
    p = P(after=2, center=True)
    R(p, C.FONO + "   \u00b7   " + C.MAIL + "   \u00b7   " + C.CIUDAD, SANS, 8.5, color=META, track=8)
    r = P(after=0, line=1)
    b = OxmlElement("w:pBdr"); bt = el("w:bottom", val="double", sz=8, space=6, color=HAIR)
    b.append(bt); r._p.get_or_add_pPr().append(b); R(r, "", SANS, 2)

    head("Perfil profesional", 12)
    for i, t in enumerate(C.PERFIL):
        q = P(before=0 if i == 0 else kp(4.5), just=True); R(q, t, SERIF, kb(9.5))

    head("\u00c1reas de expertise", 13)
    p = P(just=False, center=True, line=kl(1.3)); no_hyph(p)
    for i, t in enumerate(C.EXPERTISE):
        if i: R(p, "  \u00b7  ", SERIF, 8.5, color=ACC)
        R(p, t, SERIF, 8.5)

    def job(j, brk=False, before=kj(10)):
        p = P(before=before, keep=True, tabs=[(W, WD_TAB_ALIGNMENT.RIGHT)], brk=brk)
        R(p, j["empresa"], SERIF, kb(10.5), bold=True, color=INK); p.add_run("\t")
        R(p, j["fechas"], SANS, 7.6, bold=True, color=ACC, track=18, caps=True)
        q = P(before=1.5, after=3.5, keep=True)
        R(q, j["cargo"], SERIF, kb(9.3), color=META, italic=True)
        keep_block([p, q] + [bullet(*b) for b in j["bullets"]])

    head("Experiencia profesional", 13)
    job(C.JOBS[0], before=0); job(C.JOBS[1])
    job(C.JOBS[2], brk=True); job(C.JOBS[3])

    head("Formaci\u00f3n acad\u00e9mica", 14)
    for i, (t, inst, y) in enumerate(C.EDU):
        p = P(before=0 if i == 0 else 6.5, keep=True, tabs=[(W, WD_TAB_ALIGNMENT.RIGHT)])
        R(p, t, SERIF, kb(9.5), bold=True, color=INK); p.add_run("\t")
        R(p, y, SANS, 7.6, bold=True, color=ACC, track=18, caps=True)
        q = P(before=1); R(q, inst, SERIF, kb(8.8), color=META, italic=True)
        keep_block([p, q])

    head("Cursos y certificaciones", 13)
    for k, r_ in C.CERTS: bullet("", k, r_)

    head("Competencias", 13)
    p = P(just=False, center=True, line=kl(1.3)); no_hyph(p)
    for i, t in enumerate(C.COMPETENCIAS):
        if i: R(p, "  \u00b7  ", SERIF, 8.5, color=ACC)
        R(p, t, SERIF, 8.5)

    head("Herramientas y disponibilidad", 13)
    bullet("", "Herramientas: ", C.HERRAMIENTAS)
    bullet("", "Disponibilidad: ", C.DISPONIBILIDAD)

    f = OUT + "CV_Jeniffer_Mieres_05_Clasico.docx"; doc.save(f); return f

# =============================================================================
# VARIANTE 6 — IMPACTO (franjas de seccion en color, sans, acento pizarra)
# =============================================================================
def impacto():
    F = "Calibri"
    INK = RGBColor(0x17,0x1A,0x1E); BODY = RGBColor(0x33,0x37,0x3B)
    META = RGBColor(0x6E,0x73,0x78); ACC = RGBColor(0x24,0x43,0x5C)
    WHITE = RGBColor(0xFF,0xFF,0xFF); FILL = "24435C"
    doc = Document(); s = doc.sections[0]
    s.top_margin, s.bottom_margin = Cm(km(1.2)), Cm(km(1.1))
    s.left_margin, s.right_margin = Cm(1.6), Cm(1.5)
    W = s.page_width - s.left_margin - s.right_margin
    base_style(doc, F, kb(9.5), BODY, kl(1.16)); hyphenate(doc, zone=510)

    def P(before=0, after=0, line=None, keep=False, indent=None, hang=None, tabs=(), just=False, brk=False):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.space_before, pf.space_after = Pt(before), Pt(after)
        pf.line_spacing = kl(1.16) if line is None else line
        if keep: pf.keep_with_next = True
        if indent is not None: pf.left_indent = indent
        if hang is not None: pf.first_line_indent = hang
        for pos, al in tabs: pf.tab_stops.add_tab_stop(pos, al)
        if just: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if brk: pf.page_break_before = True
        return p

    def head(t, before=13, brk=False):
        h = P(before=ks(before), after=5, keep=True, line=1.35, brk=brk)
        shade(h, FILL)
        R(h, "  " + t.upper(), F, 8.5, bold=True, color=WHITE, track=32)

    def bullet(pre, key, post):
        p = P(after=kbu(3.6), indent=Cm(0.4), hang=Cm(-0.4), just=True)
        R(p, "\u25aa", F, 6.5, color=ACC); p.add_run("\t")
        if pre: R(p, pre, F, kb(9.5))
        if key: R(p, key, F, kb(9.5), bold=True, color=INK)
        if post: R(p, post, F, kb(9.5))
        return p

    p = P(after=2); R(p, C.NOMBRE_T, F, 26, bold=True, color=INK)
    p = P(after=2.5)
    R(p, C.TITULO.upper() + "   |   " + C.CLAIM.upper(), F, 8.5, bold=True, color=ACC, track=30)
    p = P(after=0)
    R(p, C.FONO + "   \u00b7   " + C.MAIL + "   \u00b7   " + C.CIUDAD, F, 9, color=META)
    border(p, "bottom", FILL, 22, space=6)

    head("Perfil profesional", 12)
    for i, t in enumerate(C.PERFIL):
        q = P(before=0 if i == 0 else kp(4.5), just=True); R(q, t, F, kb(9.5))

    head("\u00c1reas de expertise")
    p = P(line=kl(1.3)); no_hyph(p)
    for i, t in enumerate(C.EXPERTISE):
        if i: R(p, "  \u00b7  ", F, 9, color=RGBColor(0xAE,0xB3,0xB6))
        R(p, t, F, 9)

    def job(j, brk=False, before=kj(10.5)):
        p = P(before=before, keep=True, tabs=[(W, WD_TAB_ALIGNMENT.RIGHT)], brk=brk)
        R(p, j["empresa"], F, kb(10.5), bold=True, color=INK); p.add_run("\t")
        R(p, j["fechas"], F, 7.6, bold=True, color=ACC, track=20, caps=True)
        q = P(before=1.5, after=3.5, keep=True)
        R(q, j["cargo"], F, kb(9), color=META, italic=True)
        keep_block([p, q] + [bullet(*b) for b in j["bullets"]])

    head("Experiencia profesional")
    job(C.JOBS[0], before=0); job(C.JOBS[1])
    job(C.JOBS[2], brk=True); job(C.JOBS[3])

    head("Formaci\u00f3n acad\u00e9mica", 14)
    for i, (t, inst, y) in enumerate(C.EDU):
        p = P(before=0 if i == 0 else 6.5, keep=True, tabs=[(W, WD_TAB_ALIGNMENT.RIGHT)])
        R(p, t, F, kb(9.5), bold=True, color=INK); p.add_run("\t")
        R(p, y, F, 7.6, bold=True, color=ACC, track=20, caps=True)
        q = P(before=1); R(q, inst, F, kb(8.8), color=META, italic=True)
        keep_block([p, q])

    head("Cursos y certificaciones")
    for k, r_ in C.CERTS: bullet("", k, r_)

    head("Competencias")
    p = P(line=kl(1.3)); no_hyph(p)
    for i, t in enumerate(C.COMPETENCIAS):
        if i: R(p, "  \u00b7  ", F, 9, color=RGBColor(0xAE,0xB3,0xB6))
        R(p, t, F, 9)

    head("Herramientas y disponibilidad")
    bullet("", "Herramientas: ", C.HERRAMIENTAS)
    bullet("", "Disponibilidad: ", C.DISPONIBILIDAD)

    f = OUT + "CV_Jeniffer_Mieres_06_Impacto.docx"; doc.save(f); return f

BUILDERS = dict(swiss=swiss, banda=banda, sidebar=sidebar, editorial=rail, clasico=centro, impacto=impacto)
if __name__ == "__main__":
    for n, fn in BUILDERS.items(): print("OK ->", fn())
